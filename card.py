import hashlib
from date_test import generate_date_from_cite
from docx.text.paragraph import Paragraph
import re

Paragraph.text = property(lambda self: GetParagraphText(self))

def GetParagraphText(paragraph):
  def GetTag(element):
      return "%s:%s" % (element.prefix, re.match("{.*}(.*)", element.tag).group(1))

  text = ''
  runCount = 0
  for child in paragraph._p:
      tag = GetTag(child)
      if tag == "w:r":
          text += paragraph.runs[runCount].text
          runCount += 1
      if tag == "w:hyperlink":
          for subChild in child:
              if GetTag(subChild) == "w:r":
                  text += subChild.text
  return text

TAG_NAME = "Heading 4"
NORMAL_NAME = "Normal"
EMPHASIS_NAME = "Emphasis"
UNDERLINE_NAME = "Underline"
LIST_PARAGRAPH_NAME = "List Paragraph"
CITE_NAME = "13 pt Bold"
class Card():
  def __init__(self, paragraphs, additional_info):
    if paragraphs[0].style.name != TAG_NAME or len(paragraphs) < 2:
      raise Exception("Invalid paragraph structure")

    self.paragraphs = paragraphs
    self.tag = paragraphs[0].text.strip(", ")
    self.tag_sub = ""
    for i in range(1, len(paragraphs)):
      if not any(c.isdigit() for c in paragraphs[i].text):
        self.tag_sub += paragraphs[i].text + " "
      else:
        self.cite = paragraphs[i].text
        self.cite_i = i
        self.body = [p.text for p in paragraphs[i+1:] if p.style.name == NORMAL_NAME or p.style.name == LIST_PARAGRAPH_NAME]
        break

    if not self.body or len("".join(self.body)) < 25:
      raise Exception("Card is too short")

    self.cite_emphasis = []
    self.highlights = []
    self.highlighted_text = ""
    self.emphasis = []
    self.underlines = []
    self.parse_paragraphs()

    self.additional_info = additional_info
    self.object_id = hashlib.sha256(str(self).encode()).hexdigest()
    self.cite_date = generate_date_from_cite(self.cite)

  def parse_paragraphs(self):
    j = 0

    for r in self.paragraphs[self.cite_i].runs:
      run_text = r.text.strip()
      run_index = self.paragraphs[self.cite_i].text.find(run_text, j)

      if run_index == -1:
        continue
      if CITE_NAME in r.style.name or (r.style.font.bold or r.font.bold):
        self.cite_emphasis.append((run_index, run_index + len(run_text)))
      
      j = run_index + len(run_text)

    p_index = 2

    for i in range(self.cite_i + 1, len(self.paragraphs)):
      p = self.paragraphs[i]
      runs = p.runs
      j = 0

      for r in runs:
        run_text = r.text.strip()
        run_index = p.text.find(run_text, j)

        if run_index == -1:
          continue
        if r.font.highlight_color is not None:
          self.highlights.append((p_index, run_index, run_index + len(run_text)))
          self.highlighted_text += " " + run_text
        if UNDERLINE_NAME in r.style.name or r.font.underline or r.style.font.underline:
          self.underlines.append((p_index, run_index, run_index + len(run_text)))
        if EMPHASIS_NAME in r.style.name:
          self.emphasis.append((p_index, run_index, run_index + len(run_text)))
        
        j = run_index + len(run_text)
      
      p_index += 1
  
  def get_index(self):
    index = {
      "tag": self.tag,
      "cite": self.cite,
      "body": self.body,
      "id": self.object_id,
      "highlighted_text": self.highlighted_text,
      **self.additional_info
    }
    if self.cite_date is not None:
      index["cite_date"] = self.cite_date.strftime("%Y-%m-%d")
    return index
  
  def get_dynamo(self):
    db_representation = {
      "tag": {"S": self.tag},
      "tag_sub": {"S": self.tag_sub},
      "cite": {"S": self.cite},
      "highlighted_text": {"S": self.highlighted_text},
      "body": {"L": [{"S":p} for p in self.body]},
      "highlights": { "L": [ { "L": [ {"N":str(i)} for i in v ] } for v in self.highlights ] },
      "emphasis": {"L": [{"L": [{"N":str(i)} for i in v]} for v in self.emphasis]},
      "underlines": {"L": [{"L": [{"N":str(i)} for i in v]} for v in self.underlines]},
      "cite_emphasis": {"L": [{"L": [{"N":str(i)} for i in v]} for v in self.cite_emphasis]},
      "id": {"S": self.object_id}
    }
    
    if self.additional_info.get("division") is not None and self.additional_info.get("year") is not None and self.additional_info.get("download_url") is not None:
      db_representation["division"] = {"S": self.additional_info["division"]}
      db_representation["year"] = {"S": self.additional_info["year"]}
      db_representation["download_url"] = {"S": self.additional_info["download_url"]}
      db_representation["filename"] = {"S": self.additional_info["filename"]}
      db_representation["school"] = {"S": self.additional_info["school"]}
      db_representation["team"] = {"S": self.additional_info["team"]}
    
    if self.cite_date is not None:
      db_representation["cite_date"] = {"S": self.cite_date.strftime("%Y-%m-%d")}

    return db_representation

  def __str__(self):
    return f"{self.tag}\n{self.cite}\n{self.body}\n"

  def __repr__(self):
    return f"\n{self.tag}\n{self.cite}\n{self.body}\n"